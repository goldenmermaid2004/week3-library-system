"""Generate the Week 3 project documentation as a .docx file."""
from __future__ import annotations

import pathlib

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
OUT = ROOT / "Week3-Library-Management-System-Documentation.docx"

ACCENT = RGBColor(0x1D, 0x4E, 0x89)
DARKTEXT = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x47, 0x55, 0x69)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for level, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
    style = doc.styles[level]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = ACCENT
    style.font.bold = True


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def para(text, bold=False, italic=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)


def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = header
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            t.rows[r].cells[c].text = str(value)
    doc.add_paragraph()


def picture(filename, caption, width=6.0):
    path = DOCS / filename
    if not path.exists():
        para(f"[Screenshot pending: {filename}]", italic=True, color=GRAY)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = para("Figure: " + caption, italic=True, color=GRAY, size=9,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.paragraph_format.space_after = Pt(14)


# ============================ Title page ============================
para("Week 3 Project", color=GRAY, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
title = para("Console-Based Library Management System", bold=True, color=ACCENT,
             size=26, align=WD_ALIGN_PARAGRAPH.CENTER)
title.paragraph_format.space_before = Pt(100)
para("Java Programming Basics — OOP, Collections, File I/O & Console Applications",
     color=DARKTEXT, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Built with Java 17 — classes, ArrayLists, LocalDate, Streams, and text-file persistence",
     italic=True, color=GRAY, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Repository: https://github.com/goldenmermaid2004/week3-library-system",
     color=GRAY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ============================ 1. Project Description / Overview ============================
h1("1. Project Overview")
para("Project Description", bold=True)
para("A Java console application for managing library operations including book tracking, "
     "member management, and a borrowing system with file-based data persistence. "
     "Librarians can add and remove books, register members, borrow and return titles "
     "with due dates, calculate overdue fines, reserve books, generate statistics, and "
     "export data to CSV.")

para("Goals and objectives:", bold=True)
bullets([
    "Apply Java syntax, data types, control structures, methods, and exception handling.",
    "Practice Object-Oriented Programming: classes, objects, encapsulation, constructors, "
    "getters and setters.",
    "Use ArrayList collections to store and manage books and members.",
    "Implement file I/O so catalog and member data survive program restarts.",
    "Build a console menu with input validation and clear user feedback.",
    "Add search/filter, overdue fines, reservations, statistics, and CSV export.",
])

picture("screenshot-menu.png", "Main console menu after program start")

# ============================ 2. Features ============================
h1("2. Features")
table(["Feature", "Description"], [
    ["Add / remove books", "Manage the catalog by ISBN with duplicate and borrow checks"],
    ["Search books", "Find books by title, author, or ISBN (case-insensitive)"],
    ["Filter books", "Show available-only or borrowed-only lists"],
    ["Register members", "Store member ID, name, and email with validation"],
    ["Borrow books", "2-week loan period with due-date tracking"],
    ["Return books", "Clears borrow state and reports overdue fines if any"],
    ["Overdue fines", "$1.00 per day late using LocalDate and ChronoUnit"],
    ["Reservations", "FIFO waitlist when a book is unavailable"],
    ["Library statistics", "Totals for books, members, overdue items, and outstanding fines"],
    ["CSV export", "Exports books and members to data/books_export.csv and members_export.csv"],
    ["File persistence", "Auto load/save via data/books.txt and data/members.txt"],
    ["Input validation", "Non-empty fields, year range, email format, menu choice range"],
])

# ============================ 3. Setup Instructions ============================
h1("3. Setup Instructions")
para("Step-by-step installation and configuration:", bold=True)
numbered([
    "Install JDK 17 or later and confirm with: java -version and javac -version.",
    "Clone the repository: git clone https://github.com/goldenmermaid2004/week3-library-system.git",
    "Open a terminal in the project root: cd week3-library-system",
    "Compile: javac -d bin src/main/java/library/*.java",
    "Run from the project root (so the data/ folder is found): java -cp bin library.Main",
    "Sample books and members load automatically from data/books.txt and data/members.txt.",
])
para("Optional — Maven:", bold=True)
code_block("mvn -q package\njava -jar target/week3-library-system-1.0.0.jar")
para("Requirements: JDK 17+. No third-party libraries are required for the console application.")

# ============================ 4. How to Run ============================
h1("4. How to Run")
code_block(
    "# Compile and run\n"
    "javac -d bin src/main/java/library/*.java\n"
    "java -cp bin library.Main"
)

h2("4.1 Sample Menu")
code_block(
    "=== LIBRARY MANAGEMENT SYSTEM ===\n"
    "1.  Add New Book\n"
    "2.  View All Books\n"
    "3.  Search Books\n"
    "4.  Filter Books (Available / Borrowed)\n"
    "5.  Remove Book\n"
    "6.  Register Member\n"
    "7.  View All Members\n"
    "8.  Borrow Book\n"
    "9.  Return Book\n"
    "10. Reserve Book\n"
    "11. View Overdue Books & Fines\n"
    "12. View Library Statistics\n"
    "13. Export Data to CSV\n"
    "14. Exit\n"
    "\n"
    "Enter your choice:"
)
para("The sample brief shows a shorter menu (items 1–8). This project implements the "
     "full requirement set, so the live menu includes filter, remove, members list, "
     "reservations, overdue fines, and CSV export as additional options.")

# ============================ 5. Code Structure ============================
h1("5. Code Structure")
code_block(
    "week3-library-system/\n"
    "|-- src/main/java/library/\n"
    "|   |-- Main.java          # Console menu and input validation\n"
    "|   |-- Book.java          # Book model (due dates, fines, reservations)\n"
    "|   |-- Member.java        # Member model and borrowed ISBN list\n"
    "|   |-- Library.java       # Business logic / library operations\n"
    "|   |-- FileHandler.java   # Load/save text files + CSV export\n"
    "|-- src/main/resources/\n"
    "|-- data/\n"
    "|   |-- books.txt          # Persisted book catalog\n"
    "|   |-- members.txt        # Persisted members\n"
    "|-- docs/                 # Screenshots, smoke test, doc generators\n"
    "|-- README.md\n"
    "|-- .gitignore\n"
    "|-- pom.xml"
)
para("Class responsibilities:", bold=True)
bullets([
    "Main — presentation layer: menu loop, prompts, and validation helpers.",
    "Library — service layer: add/remove/search, borrow/return, reservations, stats, export.",
    "Book / Member — domain models with private fields and public getters/setters.",
    "FileHandler — persistence layer: pipe-delimited text files and CSV writers.",
])

# ============================ 6. Technical Details ============================
h1("6. Technical Details")

h2("6.1 Architecture")
para("The application follows a simple layered design:")
code_block(
    "User (console)\n"
    "   -> Main (menu / validation)\n"
    "      -> Library (business rules)\n"
    "         -> Book / Member (domain state)\n"
    "         -> FileHandler (books.txt / members.txt / CSV)"
)

h2("6.2 Data structures")
bullets([
    "ArrayList<Book> — in-memory catalog.",
    "ArrayList<Member> — registered members.",
    "List<String> on Member — ISBNs currently borrowed.",
    "List<String> on Book — FIFO reservation queue of member IDs.",
    "Java Streams — filter, findFirst, collect, count, mapToDouble for search and stats.",
])

h2("6.3 Algorithms")
bullets([
    "Borrow: validate book/member existence and availability, enforce borrow limit (5), "
    "honor reservation head-of-queue, set due date to LocalDate.now().plusWeeks(2), persist.",
    "Return: verify borrower, compute fine if overdue, clear borrow fields, notify next reservation.",
    "Overdue fine: ChronoUnit.DAYS.between(dueDate, today) × $1.00.",
    "Search: case-insensitive contains match on title, author, or ISBN.",
    "Persistence: serialize each record as pipe-separated fields; use '-' for empty optionals.",
])

h2("6.4 File formats")
code_block(
    "Book line:\n"
    "ISBN|title|author|year|available|borrowedBy|dueDate|reservations\n"
    "\n"
    "Member line:\n"
    "id|name|email|borrowedIsbn1,borrowedIsbn2\n"
    "\n"
    "Empty optional fields use the token '-'."
)

h2("6.5 Exception handling and validation")
bullets([
    "FileHandler wraps read/write in try/catch and prints clear error messages instead of crashing.",
    "Malformed lines are skipped with a console warning.",
    "Main rejects empty strings, invalid years, weak emails, and out-of-range menu choices.",
    "Business rules reject duplicate ISBNs/IDs, removing borrowed books, and illegal borrows.",
])

picture("screenshot-validation.png", "Input validation and user feedback examples")

# ============================ 7. Visual Documentation ============================
h1("7. Visual Documentation")
para("The following screenshots were generated from live console sessions of the compiled program.")
picture("screenshot-books.png", "View All Books — sample catalog with availability and overdue status")
picture("screenshot-search.png", "Search Books — keyword 'Java' finds Effective Java")
picture("screenshot-stats.png", "Library statistics summary")
picture("screenshot-overdue.png", "Overdue books with calculated fines")

# ============================ 8. Testing Evidence ============================
h1("8. Testing Evidence")
para("Manual and scripted tests were run against the compiled application. "
     "A smoke-test input file (docs/smoke_input.txt) exercises view, search, statistics, "
     "and overdue listing non-interactively.")
code_block(
    "# Windows PowerShell (from project root)\n"
    "javac -d bin src/main/java/library/*.java\n"
    "Get-Content docs/smoke_input.txt | java -cp bin library.Main"
)

table(["#", "Test case", "Expected result", "Status"], [
    [1, "Start program", "Loads 5 books and 5 members; menu displayed", "PASS"],
    [2, "View all books", "Formatted list of 5 sample books", "PASS"],
    [3, "Search 'Java'", "Matches Effective Java", "PASS"],
    [4, "View statistics", "Totals match catalog (e.g. 5 books, 5 members)", "PASS"],
    [5, "View overdue books", "Spring in Action shows fine ($1/day late)", "PASS"],
    [6, "Empty ISBN on add", "Rejected: input cannot be empty", "PASS"],
    [7, "Duplicate ISBN", "Rejected: book already exists", "PASS"],
    [8, "Duplicate member ID", "Rejected: member already exists", "PASS"],
    [9, "Borrow missing ISBN", "Book not found message", "PASS"],
    [10, "Invalid menu text", "Prompts again for a number", "PASS"],
    [11, "Borrow available book", "Due date ~2 weeks ahead; file updated", "PASS"],
    [12, "Return on-time book", "No fine message", "PASS"],
    [13, "Return overdue book", "Fine printed and book becomes available", "PASS"],
    [14, "Reserve borrowed book", "Queue position shown; saved to books.txt", "PASS"],
    [15, "Remove borrowed book", "Rejected with clear error", "PASS"],
    [16, "Export CSV", "Creates data/books_export.csv and members_export.csv", "PASS"],
    [17, "Restart program", "Previous changes reloaded from text files", "PASS"],
])
para("Smoke-test scripted run: menu, book list, search, statistics, and overdue views "
     "completed successfully with exit code 0.", italic=True, color=GRAY)

# ============================ 9. Quality Standards Checklist ============================
h1("9. Quality Standards Checklist")
para("All items required for full marks:", bold=True)
table(["Requirement", "Included"], [
    ["Project Overview — clear goals and objectives", "Yes — Section 1"],
    ["Setup Instructions — step-by-step install/config", "Yes — Section 3"],
    ["Code Structure — organized file hierarchy", "Yes — Section 5"],
    ["Visual Documentation — screenshots of functionality", "Yes — Sections 1, 6, 7"],
    ["Technical Details — algorithms, structures, architecture", "Yes — Section 6"],
    ["Testing Evidence — test cases and validation", "Yes — Section 8"],
    ["Features list + How to Run + Sample Menu", "Yes — Sections 2 and 4"],
    ["OOP encapsulation, ArrayLists, File I/O, exceptions", "Yes — implemented in source"],
])

# ============================ 10. How to use ============================
h1("10. How to Use (Quick Guide)")
numbered([
    "Start the program and choose an option from the numbered menu.",
    "Add books (option 1) or browse/search/filter the catalog (options 2–4).",
    "Register members (option 6) before borrowing.",
    "Borrow with ISBN + Member ID (option 8); return with the same pair (option 9).",
    "Check overdue fines (option 11) and overall statistics (option 12).",
    "Reserve unavailable titles (option 10) or export CSV backups (option 13).",
    "Choose Exit (option 14) when finished — all changes are already saved to disk.",
])

doc.save(str(OUT))
print("Saved:", OUT)
